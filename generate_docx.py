import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = docx.Document()

# Title
title = doc.add_heading('Project ALO: Detailed System Architecture & AI Vision', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 1. Executive Summary
doc.add_heading('1. Executive Summary', level=1)
doc.add_paragraph(
    "Project ALO is Airtel's next-generation network intelligence platform. Moving beyond static reporting, this platform provides real-time geospatial visualization, hierarchical network health monitoring (from Circles to individual MSANs), and proactive alerting."
)
doc.add_paragraph(
    "In its enhanced architecture, the platform utilizes Google Spreadsheets as a dynamic, centralized Data Backend, allowing network engineers and on-ground staff to update network metrics without requiring database redeployments. Furthermore, the platform serves as the foundational layer for integrating advanced Artificial Intelligence (A.I.) for predictive maintenance and automated network optimization."
)

# 2. Comprehensive System Architecture
doc.add_heading('2. Comprehensive System Architecture', level=1)
doc.add_paragraph("The architecture is designed to be serverless, highly responsive, and easily updatable via Google Workspace.")

doc.add_heading('Component Details', level=2)
p = doc.add_paragraph()
p.add_run('Google Spreadsheet Backend: ').bold = True
p.add_run('Acts as the Headless CMS/Database. Different sheets represent different entities (e.g., Circles, OLT_Cities, BRAS_Metrics, Live_Complaints).')
p.style = 'List Bullet'

p = doc.add_paragraph()
p.add_run('Next.js API Route: ').bold = True
p.add_run('A server-side function that securely fetches data from the Google Sheets API using a Service Account, caching the response to prevent API rate limiting.')
p.style = 'List Bullet'

p = doc.add_paragraph()
p.add_run('Data Transformation Service: ').bold = True
p.add_run('Converts the flat 2D array data from Google Sheets into the hierarchical JSON structure (Circle -> Hub -> OLT -> BRAS -> MSAN) required by the frontend.')
p.style = 'List Bullet'

p = doc.add_paragraph()
p.add_run('Dashboard Context: ').bold = True
p.add_run('React\'s global state manager that distributes the transformed data to all visual components instantly.')
p.style = 'List Bullet'

# 3. Dynamic Data Flow & Alerting
doc.add_heading('3. Dynamic Data Flow & Alerting', level=1)
doc.add_paragraph("Lifecycle of a network complaint update, from the moment an engineer updates the Google Sheet to when the AI processes it and the Dashboard visualizes it:")

doc.add_paragraph("1. NOC Engineer updates complaint count for OLT 'Hyderabad' in Google Sheets.", style='List Number')
doc.add_paragraph("2. Google Sheet auto-saves data.", style='List Number')
doc.add_paragraph("3. Dashboard Client periodically polls Next.js API for updates.", style='List Number')
doc.add_paragraph("4. Next.js API checks Edge Cache validity. If expired, it securely fetches latest rows via Google Sheets API.", style='List Number')
doc.add_paragraph("5. Next.js API transforms the 2D array to Hierarchical JSON, updates cache, and returns it to the Client.", style='List Number')
doc.add_paragraph("6. Client UI recalculates Severity (Complaint-to-Subscriber Ratio), re-ranks the Alerts Sidebar, and updates Map Node colors (e.g., Red for Critical).", style='List Number')
doc.add_paragraph("7. In parallel, the future AI Engine ingests the historical data, runs anomaly detection, and pushes predictive failure alerts to the Dashboard.", style='List Number')

# 4. Management Vision: AI-Driven AIOps Roadmap
doc.add_heading('4. Management Vision: AI-Driven AIOps Roadmap', level=1)
doc.add_paragraph("To transition from a reactive monitoring tool to a Proactive Network Intelligence Platform, Project ALO will integrate advanced A.I. processing directly connected to the Google Sheets data pipeline.")

doc.add_heading('Phase 1: Predictive Analytics (Months 1-3)', level=2)
p = doc.add_paragraph()
p.add_run('Data Ingestion: ').bold = True
p.add_run('A python-based microservice will continuously read historical MSAN/BRAS metrics and complaint logs from the Google Sheet.')
p.style = 'List Bullet'
p = doc.add_paragraph()
p.add_run('Failure Prediction: ').bold = True
p.add_run('Machine learning models will identify patterns preceding equipment degradation.')
p.style = 'List Bullet'
p = doc.add_paragraph()
p.add_run('Actionable Output: ').bold = True
p.add_run('The AI will write a new Risk_Score back to the Google Sheet. The dashboard will instantly visualize high-risk nodes in Orange/Yellow.')
p.style = 'List Bullet'

doc.add_heading('Phase 2: Intelligent Capacity Planning (Months 3-6)', level=2)
p = doc.add_paragraph()
p.add_run('Load Forecasting: ').bold = True
p.add_run('AI models will correlate subscriber growth trends from the sheet with network congestion in specific OLT cities.')
p.style = 'List Bullet'
p = doc.add_paragraph()
p.add_run('Capex Optimization: ').bold = True
p.add_run('The dashboard will feature an Investment Insights panel, highlighting OLT cities requiring bandwidth reallocation for maximum ROI.')
p.style = 'List Bullet'

doc.add_heading('Phase 3: Automated Remediation (Months 6-12)', level=2)
p = doc.add_paragraph()
p.add_run('Root Cause Analysis (RCA): ').bold = True
p.add_run('When a sudden spike in complaints is registered, the AI will analyze logs across multiple network layers to pinpoint fiber cuts, power failures, or software issues.')
p.style = 'List Bullet'
p = doc.add_paragraph()
p.add_run('Smart Triaging: ').bold = True
p.add_run('The system will automatically dispatch the correct team (e.g., Fiber Team vs. Power Team) via SMS/Email directly from the platform.')
p.style = 'List Bullet'

# 5. Security & Implementation Details for Google Sheets
doc.add_heading('5. Security & Implementation Details', level=1)
p = doc.add_paragraph()
p.add_run('Authentication: ').bold = True
p.add_run('The Next.js backend will authenticate with Google using an IAM Service Account (JSON key) rather than user-level OAuth, ensuring uninterrupted background access.')
p.style = 'List Bullet'
p = doc.add_paragraph()
p.add_run('Data Validation: ').bold = True
p.add_run('Strict Zod validations will ensure that if a NOC engineer accidentally types text into a number column, the system gracefully falls back to previous safe data.')
p.style = 'List Bullet'
p = doc.add_paragraph()
p.add_run('Rate Limiting Mitigation: ').bold = True
p.add_run('Google Sheets API rate limits will be avoided using SWR (Stale-While-Revalidate) and Edge Caching, ensuring the spreadsheet is queried safely regardless of simultaneous users.')
p.style = 'List Bullet'

doc.save('Detailed_Project_ALO_Architecture.docx')
print("Detailed Word document created successfully.")
