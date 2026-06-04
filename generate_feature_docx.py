import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image, ImageDraw, ImageFont

OUTPUT_DOCX = 'Airtel_Dashboard_Feature_Summary.docx'
DIAGRAM1_PNG = 'feature_architecture.png'
DIAGRAM2_PNG = 'feature_flow.png'


def make_diagram_1(path: str) -> None:
    width, height = 1200, 720
    bg = (17, 24, 39)
    fg = (226, 232, 240)
    accent = (59, 130, 246)
    accent2 = (168, 85, 247)
    img = Image.new('RGB', (width, height), bg)
    draw = ImageDraw.Draw(img)
    font = ImageFont.load_default()

    boxes = [
        ((80, 120, 420, 260), 'Data Layer\nGoogle Sheets'),
        ((430, 80, 830, 300), 'Backend Layer\nNext.js API'),
        ((840, 120, 1160, 260), 'Frontend Layer\nReact UI'),
    ]
    for box, text in boxes:
        draw.rectangle(box, outline=fg, width=3)
        draw.multiline_text((box[0] + 20, box[1] + 20), text, fill=fg, font=font, spacing=4)

    arrow = [(420, 190), (430, 190), (420, 185), (420, 195)]
    draw.line([(420, 190), (430, 190)], fill=accent, width=4)
    draw.polygon(arrow, fill=accent)
    arrow = [(830, 190), (840, 190), (830, 185), (830, 195)]
    draw.line([(830, 190), (840, 190)], fill=accent2, width=4)
    draw.polygon(arrow, fill=accent2)

    lines = [
        ((250, 260), (250, 520)),
        ((250, 520), (180, 520)),
        ((180, 520), (180, 610)),
    ]
    for line in lines:
        draw.line(line, fill=fg, width=3)
    draw.text((380, 340), 'Transform 2D sheet rows into\nhierarchical JSON payload', fill=fg, font=font)

    draw.text((140, 540), 'Sheet data → API cache → UI context', fill=fg, font=font)
    img.save(path)


def make_diagram_2(path: str) -> None:
    width, height = 1200, 520
    bg = (15, 23, 42)
    fg = (226, 232, 240)
    accent = (59, 130, 246)
    img = Image.new('RGB', (width, height), bg)
    draw = ImageDraw.Draw(img)
    font = ImageFont.load_default()

    steps = [
        ('1. Update Google Sheet', 120, 100),
        ('2. Poll Next.js API', 120, 230),
        ('3. Fetch & Cache Data', 120, 360),
    ]
    for label, x, y in steps:
        draw.rectangle((x, y, x+360, y+90), outline=fg, width=3)
        draw.text((x+20, y+20), label, fill=fg, font=font)
    draw.line((300, 190, 300, 230), fill=accent, width=4)
    draw.line((300, 320, 300, 360), fill=accent, width=4)
    draw.polygon([(295, 230), (305, 230), (300, 240)], fill=accent)
    draw.polygon([(295, 360), (305, 360), (300, 370)], fill=accent)
    draw.text((560, 150), '4. Transform & Validate', fill=fg, font=font)
    draw.text((560, 180), '5. Return hierarchical JSON', fill=fg, font=font)
    draw.text((560, 210), '6. Re-rank alerts & redraw map', fill=fg, font=font)
    img.save(path)


def build_doc() -> None:
    doc = Document()
    title = doc.add_heading('Airtel Network Dashboard — Feature Summary', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('A compact overview of the current dashboard features, component architecture, and interaction flow.')

    doc.add_heading('1. Key Features', level=1)
    features = [
        ('Network Map', 'Interactive India map with OLT, complaint, and ideal topologies.'),
        ('Alerts Sidebar', 'Live-ranked complaints and trouble spots by city and ratio.'),
        ('BNG Utilization', 'Hierarchical drilldown from circles to AE interfaces with utilization filtering.'),
        ('Mobile Networks', 'District-level 4G/5G volume visualization with trends and map drilldown.'),
        ('Mobile Hubs', 'Hub-to-district connections, volume analytics, and circle filters.'),
        ('Top Controls', 'Date slider, 4G/5G mode, line toggle, fixed coordinate highlights.'),
    ]
    for label, text in features:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{label}: ').bold = True
        p.add_run(text)

    doc.add_heading('2. Architecture Overview', level=1)
    doc.add_paragraph('The dashboard uses Google Sheets as a data backend, a Next.js API layer for transformation and caching, and a React client with Leaflet and Recharts for visualization.')
    if os.path.exists(DIAGRAM1_PNG):
        doc.add_picture(DIAGRAM1_PNG, width=Inches(6.5))

    doc.add_heading('3. Data Flow', level=1)
    doc.add_paragraph('Changes from network operations flow through the Google Sheet into the API, then to the dashboard UI, where alerts are re-ranked and visuals are updated.')
    if os.path.exists(DIAGRAM2_PNG):
        doc.add_picture(DIAGRAM2_PNG, width=Inches(6.5))

    doc.add_heading('4. User Interaction Summary', level=1)
    bullets = [
        'Select circle filters to show or hide hub groups on the map.',
        'Click hubs or districts to open detailed side panels with volume charts.',
        'Use the date slider to examine historical traffic across 4G/5G waves.',
        'Open the alerts panel to review complaint ratios and prioritize action.',
    ]
    for item in bullets:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('5. Notes', level=1)
    doc.add_paragraph('This document is generated from the current dashboard codebase and includes visuals to explain the platform flow and features. It is suitable for stakeholder review or handoff documentation.')

    doc.save(OUTPUT_DOCX)
    print(f'Created {OUTPUT_DOCX}')


if __name__ == '__main__':
    make_diagram_1(DIAGRAM1_PNG)
    make_diagram_2(DIAGRAM2_PNG)
    build_doc()
