import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = docx.Document()

# Title
title = doc.add_heading('Project ALO: Detailed System Architecture', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 1. Executive Summary
doc.add_heading('1. Executive Summary', level=1)
doc.add_paragraph(
    "Project ALO is Airtel's next-generation network intelligence platform. Moving beyond static reporting, this platform provides real-time geospatial visualization, hierarchical network health monitoring (from Circles to individual MSANs), and proactive alerting."
)
doc.add_paragraph(
    "In its enhanced architecture, the platform utilizes Google Spreadsheets as a dynamic, centralized Data Backend, allowing network engineers and on-ground staff to update network metrics without requiring database redeployments. The platform translates this raw data into actionable insights for rapid network triage."
)

# 2. Comprehensive System Architecture
doc.add_heading('2. Comprehensive System Architecture', level=1)
doc.add_paragraph("The architecture is designed to be serverless, highly responsive, and easily updatable via Google Workspace.")

ascii_diagram_1 = """
+-----------------------+        +--------------------------+        +-----------------------+
| Google Workspace      |        | Next.js Serverless API   |        | React Client UI       |
|                       |        |                          |        |                       |
| [ Spreadsheets ] <----+--------+-> [ API Routes ] <-------+--------+-> [ Context API ]     |
|                       |        |         |                |        |         |             |
| [ Sheets API v4]      |        |         v                |        |         v             |
+-----------------------+        | [ Data Transformer ]     |        | [ Map & Alerts ]      |
                                 +--------------------------+        +-----------------------+
"""
p = doc.add_paragraph()
run = p.add_run(ascii_diagram_1)
run.font.name = 'Courier New'
run.font.size = Pt(9)

doc.add_heading('Component Details', level=2)
p = doc.add_paragraph()
p.add_run('Google Spreadsheet Backend: ').bold = True
p.add_run('Acts as the Headless CMS/Database. Different sheets represent different entities (e.g., Circles, OLT_Cities, BRAS_Metrics, Live_Complaints).')
p.style = 'List Bullet'

p = doc.add_paragraph()
p.add_run('Next.js API Route: ').bold = True
p.add_run('A server-side function that securely fetches data from the Google Sheets API using a Service Account, caching the response to prevent API rate limiting.')
p.style = 'List Bullet'

p = doc.add_paragraph()
p.add_run('Data Transformation Service: ').bold = True
p.add_run('Converts the flat 2D array data from Google Sheets into the hierarchical JSON structure (Circle -> Hub -> OLT -> BRAS -> MSAN) required by the frontend.')
p.style = 'List Bullet'

p = doc.add_paragraph()
p.add_run('Dashboard Context: ').bold = True
p.add_run('React\'s global state manager that distributes the transformed data to all visual components instantly.')
p.style = 'List Bullet'

# 3. Dynamic Data Flow & Alerting
doc.add_heading('3. Dynamic Data Flow & Alerting', level=1)
doc.add_paragraph("Lifecycle of a network complaint update, from the moment an engineer updates the Google Sheet to when the Dashboard visualizes it:")

ascii_diagram_2 = """
[1] Engineer Updates Google Sheet
          |
          v
[2] Sheet Auto-Saves
          |
          v
[3] UI Polls API for Refresh Request
          |
          v
[4] API Fetches Latest Rows from Google Sheets
          |
          v
[5] API Transforms 2D Array to Hierarchical JSON
          |
          v
[6] UI Recalculates Severity (Ranks Alerts, Colors Map)
"""
p = doc.add_paragraph()
run = p.add_run(ascii_diagram_2)
run.font.name = 'Courier New'
run.font.size = Pt(10)

doc.add_paragraph("1. NOC Engineer updates the complaint count for OLT 'Hyderabad' in the designated Google Sheet.", style='List Number')
doc.add_paragraph("2. Google Sheet auto-saves the data in the cloud.", style='List Number')
doc.add_paragraph("3. The Dashboard Client periodically polls the Next.js API for updates.", style='List Number')
doc.add_paragraph("4. The Next.js API checks the Edge Cache validity. If expired, it securely fetches the latest rows via the Google Sheets API.", style='List Number')
doc.add_paragraph("5. The Next.js API transforms the 2D array to Hierarchical JSON, updates the cache, and returns the payload to the Client.", style='List Number')
doc.add_paragraph("6. The Client UI recalculates the Severity Score (Complaint-to-Subscriber Ratio), re-ranks the Alerts Sidebar, and updates Map Node colors (e.g., turning a node Red for Critical).", style='List Number')

# 4. Security & Implementation Details for Google Sheets
doc.add_heading('4. Security & Implementation Details', level=1)
p = doc.add_paragraph()
p.add_run('Authentication: ').bold = True
p.add_run('The Next.js backend authenticates with Google using an IAM Service Account (JSON key) rather than user-level OAuth, ensuring secure and uninterrupted background access.')
p.style = 'List Bullet'
p = doc.add_paragraph()
p.add_run('Data Validation: ').bold = True
p.add_run('Strict Zod validations ensure that if a NOC engineer accidentally types text into a number column, the system gracefully falls back to previous safe data, preventing dashboard crashes.')
p.style = 'List Bullet'
p = doc.add_paragraph()
p.add_run('Rate Limiting Mitigation: ').bold = True
p.add_run('Google Sheets API rate limits are avoided using SWR (Stale-While-Revalidate) and Edge Caching, ensuring the spreadsheet is queried efficiently regardless of how many users are viewing the dashboard simultaneously.')
p.style = 'List Bullet'

doc.save('Project_ALO_Architecture_Final_No_AI.docx')
print("Detailed Word document with ASCII diagrams but no AI created successfully.")
