import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_DOCX = 'Airtel_Dashboard_Modes_Screens_Summary.docx'

SUMMARY = [
    ('Dashboard Modes and Screens',
     'A concise summary of the dashboard screens, interactive map modes, and the hub sidebar behavior.'),
    ('OLT Network Map',
     'The main map screen shows OLT network traffic across India with three modes: OLT Network, Complaint Map, and Ideal View.'),
    ('Complaint Map',
     'Complaint mode highlights cities by complaint volume and ratio, surfacing top trouble spots and enabling fast prioritization.'),
    ('Ideal View',
     'Ideal mode evaluates alternative BNG locations and highlights cities where a closer or more optimal BNG hub exists.'),
    ('Hub Sidebar',
     'Selecting a hub opens a right-side panel showing hub metrics, 4G/5G volume, connected districts, backhaul capacity, sparklines, and district assignment details.'),
    ('BNG Utilization Screen',
     'Shows hierarchical BNG analytics from circles to cities, BRAS nodes, and AE interfaces with filters and search, enabling drilldown into utilization and network health.'),
    ('Mobile Network Screen',
     'Displays district-level 4G/5G volume on a map with trend charts and per-district drilldown, helping compare mobile traffic across regions.'),
    ('Mobile Hubs Screen',
     'Visualizes hub-to-district connections, circle filters, and hub selection, with line visibility, fixed-coordinate highlights, and hub volume analytics.'),
]

INTERACTIONS = [
    'Top map toggle switches between OLT Network, Complaint Map, and Ideal View.',
    'Select a hub marker to open the hub sidebar or click a district marker to inspect assigned district details.',
    'Use the date slider to move through historical 4G/5G traffic and compare volume across dates.',
    'Circle filters control which hub groups are visible on the Mobile Hubs map.',
    'Complaint mode uses color and size to make problem cities immediately visible.',
    'Ideal mode surfaces routing improvements and shorter BNG alternatives for optimization.',
]


def build_doc() -> None:
    doc = Document()
    title = doc.add_heading('Airtel Dashboard Modes and Screens Summary', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('This document captures the dashboard screens, map modes, and hub sidebar interactions used in the Airtel network visualization app.')

    doc.add_heading('1. Core Screens', level=1)
    for label, text in SUMMARY[1:5]:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{label}: ').bold = True
        p.add_run(text)

    doc.add_heading('2. Hub Sidebar Behavior', level=1)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Hub Sidebar: ').bold = True
    p.add_run('When a hub is selected, a right-hand panel appears with hub KPIs, 4G/5G volume, connected district list, backhaul capacity, coordinates, and quick drilldown controls.')

    doc.add_heading('3. Detailed Screens', level=1)
    for label, text in SUMMARY[5:]:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{label}: ').bold = True
        p.add_run(text)

    doc.add_heading('4. Additional Interaction Notes', level=1)
    for item in INTERACTIONS:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph('The document is generated directly from the current app context and is ready for stakeholder handoff or quick review.')
    doc.save(OUTPUT_DOCX)
    print(f'Created {OUTPUT_DOCX}')


if __name__ == '__main__':
    build_doc()
